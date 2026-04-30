"""DOCX parser (S195).

Like ``parsers_pdf.py``, this keeps the optional ``python-docx`` dep
behind a backend Protocol so the kb-engine core stays stdlib + pydantic.

Tables become GitHub-flavoured markdown blocks (header row + separator
row + body rows). Styled headings are mapped from the docx ``Heading N``
style names to ``#`` prefixes (1\u20136). Paragraphs without a heading
style flow as prose.
"""

from __future__ import annotations

from dataclasses import dataclass
from enum import StrEnum
from typing import Protocol, runtime_checkable
from uuid import UUID

from loop_kb_engine.models import Document
from loop_kb_engine.parsers import DocumentParseError


class DocxBlockKind(StrEnum):
    PARAGRAPH = "paragraph"
    HEADING = "heading"
    TABLE = "table"


@dataclass(frozen=True, slots=True)
class DocxParagraph:
    kind = DocxBlockKind.PARAGRAPH
    text: str


@dataclass(frozen=True, slots=True)
class DocxHeading:
    kind = DocxBlockKind.HEADING
    level: int
    text: str

    def __post_init__(self) -> None:
        if not 1 <= self.level <= 6:
            raise ValueError(f"heading level must be 1..6, got {self.level}")


@dataclass(frozen=True, slots=True)
class DocxTable:
    kind = DocxBlockKind.TABLE
    rows: tuple[tuple[str, ...], ...]

    def __post_init__(self) -> None:
        if not self.rows:
            raise ValueError("table must have at least one row")
        width = len(self.rows[0])
        if width == 0:
            raise ValueError("table rows must have at least one column")
        for row in self.rows:
            if len(row) != width:
                raise ValueError("all table rows must have equal column count")


DocxBlock = DocxParagraph | DocxHeading | DocxTable


@runtime_checkable
class DocxBackend(Protocol):
    def extract(self, data: bytes) -> tuple[DocxBlock, ...]: ...


class _PythonDocxBackend:
    """Lazy ``python-docx`` backend."""

    _STYLE_TO_LEVEL: dict[str, int] = {  # noqa: RUF012
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Heading 5": 5,
        "Heading 6": 6,
        "Title": 1,
    }

    def extract(self, data: bytes) -> tuple[DocxBlock, ...]:
        try:
            from io import BytesIO

            from docx import Document as _Doc  # type: ignore[import-not-found]
        except ImportError as exc:
            raise DocumentParseError(
                "python-docx is required for DOCX parsing; install loop-kb-engine[docx]"
            ) from exc
        try:
            doc = _Doc(BytesIO(data))
        except Exception as exc:  # python-docx wraps zipfile errors broadly
            raise DocumentParseError(f"docx is unreadable: {exc}") from exc
        out: list[DocxBlock] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            level = self._STYLE_TO_LEVEL.get(para.style.name)
            if level is not None:
                out.append(DocxHeading(level=level, text=text))
            else:
                out.append(DocxParagraph(text=text))
        for tbl in doc.tables:
            rows: list[tuple[str, ...]] = []
            for row in tbl.rows:
                rows.append(tuple(cell.text.strip() for cell in row.cells))
            if rows:
                out.append(DocxTable(rows=tuple(rows)))
        return tuple(out)


def render_blocks(blocks: tuple[DocxBlock, ...]) -> str:
    """Render an extracted block list as deterministic markdown."""

    parts: list[str] = []
    for block in blocks:
        if block.kind is DocxBlockKind.HEADING:
            assert isinstance(block, DocxHeading)
            parts.append(f"{'#' * block.level} {block.text}")
        elif block.kind is DocxBlockKind.PARAGRAPH:
            assert isinstance(block, DocxParagraph)
            parts.append(block.text)
        else:  # TABLE
            assert isinstance(block, DocxTable)
            parts.append(_render_table(block))
    return "\n\n".join(parts).strip()


def _render_table(table: DocxTable) -> str:
    width = len(table.rows[0])
    header = "| " + " | ".join(table.rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * width) + " |"
    body = ["| " + " | ".join(r) + " |" for r in table.rows[1:]]
    return "\n".join([header, sep, *body])


class DocxParser:
    """``Parser`` impl for ``.docx`` documents."""

    name = "docx"

    def __init__(self, *, backend: DocxBackend | None = None) -> None:
        self._backend: DocxBackend = backend or _PythonDocxBackend()

    def parse(
        self,
        data: bytes,
        *,
        workspace_id: UUID,
        title: str,
        source: str = "",
    ) -> Document:
        if not data:
            raise DocumentParseError("docx bytes are empty")
        blocks = self._backend.extract(data)
        return self.parse_blocks(
            blocks,
            workspace_id=workspace_id,
            title=title,
            source=source,
        )

    def parse_blocks(
        self,
        blocks: tuple[DocxBlock, ...],
        *,
        workspace_id: UUID,
        title: str,
        source: str = "",
    ) -> Document:
        meta: dict[str, str] = {
            "block_count": str(len(blocks)),
            "table_count": str(sum(1 for b in blocks if b.kind is DocxBlockKind.TABLE)),
            "heading_count": str(sum(1 for b in blocks if b.kind is DocxBlockKind.HEADING)),
        }
        return Document(
            workspace_id=workspace_id,
            title=title,
            text=render_blocks(blocks),
            source=source,
            metadata=meta,
        )


__all__ = [
    "DocxBackend",
    "DocxBlock",
    "DocxBlockKind",
    "DocxHeading",
    "DocxParagraph",
    "DocxParser",
    "DocxTable",
    "render_blocks",
]
